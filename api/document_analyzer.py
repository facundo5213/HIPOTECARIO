import io
import re
import unicodedata
from datetime import date, datetime, timedelta
from decimal import Decimal, InvalidOperation
from difflib import SequenceMatcher
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

REGULATION_PATTERN = re.compile(r"\b(802\s*[_-]?\s*01|802|800)\b", re.IGNORECASE)
PRODUCT_PATTERN = re.compile(r"\bVK\d{3,5}\b", re.IGNORECASE)
CUIT_PATTERN = re.compile(r"\b\d{11}\b")
DATE_PATTERN = re.compile(r"\b(\d{1,2}/\d{1,2}/\d{4})\b")
RATE_PATTERN = re.compile(r"(?<!\d)(\d{1,2}(?:[.,]\d{1,4})?)\s*%")

def clean(value) -> str:
   if value is None:
       return ""
   return re.sub(r"\s+", " ", str(value)).strip()

def normalized(value) -> str:
   value = unicodedata.normalize("NFKD", clean(value)).encode("ascii", "ignore").decode("ascii")
   return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()

def parse_date(value):
   if isinstance(value, datetime):
       return value.date()
   if isinstance(value, date):
       return value
   match = DATE_PATTERN.search(clean(value))
   if not match:
       return None
   try:
       return datetime.strptime(match.group(1), "%d/%m/%Y").date()
   except ValueError:
       return None

def parse_rate(value):
   match = RATE_PATTERN.search(clean(value))
   if not match:
       return None
   try:
       return Decimal(match.group(1).replace(",", "."))
   except InvalidOperation:
       return None

def json_value(value):
   if isinstance(value, (date, datetime)):
       return value.isoformat()
   if isinstance(value, Decimal):
       return float(value)
   return value

def finding(kind, location, evidence, confidence, extracted, proposal_type=None, proposed=None, key=None, reason=None):
   proposed = proposed or None
   confidence = round(float(confidence), 2)
   automatic = proposal_type is not None and proposed is not None and confidence >= 90
   return {
       "type": kind,
       "location": location,
       "evidence": clean(evidence)[:1500],
       "confidence": confidence,
       "status": "PROPUESTO" if automatic else "REQUIERE_REVISION",
       "proposal_type": proposal_type,
       "proposal_key": key,
       "extracted": {k: json_value(v) for k, v in extracted.items()},
       "proposed": {k: json_value(v) for k, v in proposed.items()} if proposed else None,
       "reason": reason,
       "selected": automatic and confidence >= 95,
   }

def extract_docx(content: bytes):
   document = Document(io.BytesIO(content))
   paragraphs = [clean(item.text) for item in document.paragraphs if clean(item.text)]
   tables = []
   for table_index, table in enumerate(document.tables, start=1):
       rows = [[clean(cell.text) for cell in row.cells] for row in table.rows]
       tables.append({"location": f"Tabla {table_index}", "rows": rows})
   return {"paragraphs": paragraphs, "tables": tables, "warnings": []}

def extract_pdf(content: bytes):
   reader = PdfReader(io.BytesIO(content))
   paragraphs, warnings = [], []
   for page_number, page in enumerate(reader.pages, start=1):
       text = page.extract_text() or ""
       paragraphs.extend([f"[Página {page_number}] {clean(line)}" for line in text.splitlines() if clean(line)])
   if sum(len(item) for item in paragraphs) < 100:
       warnings.append("El PDF parece escaneado o no contiene texto extraíble. Requiere OCR o el archivo Word original.")
   return {"paragraphs": paragraphs, "tables": [], "warnings": warnings}

def extract_xlsx(content: bytes):
   workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
   tables = []
   for sheet in workbook.worksheets:
       rows = []
       for row in sheet.iter_rows(values_only=True):
           values = [clean(value) for value in row]
           if any(values):
               rows.append(values)
           if len(rows) >= 5001:
               break
       if rows:
           tables.append({"location": f"Hoja {sheet.title}", "rows": rows})
   return {"paragraphs": [], "tables": tables, "warnings": []}

def detect_regulation(text: str):
   matches = [match.group(1).replace(" ", "").replace("-", "_") for match in REGULATION_PATTERN.finditer(text)]
   normalized_matches = ["802_01" if "01" in item else item for item in matches]
   if not normalized_matches:
       return None, 0
   # La variante 802_01 es más específica; una vez mencionada explícitamente no debe
   # quedar eclipsada por referencias internas repetidas a la línea general 802.
   if "802_01" in normalized_matches:
       return "802_01", min(99, 92 + normalized_matches.count("802_01") * 2)
   counts = {code: normalized_matches.count(code) for code in set(normalized_matches)}
   code = max(counts, key=counts.get)
   return code, min(99, 88 + counts[code] * 2)

REGULATION_SHEET_NAMES = {"800", "802", "802_01"}

def sheet_regulation(location: str):
   name = location.replace("Hoja ", "").strip().upper().replace("-", "_")
   return name if name in REGULATION_SHEET_NAMES else None

def update_sheet_columns(headers):
   values = [normalized(value) for value in headers]
   def locate(include, exclude=()):
       return next((index for index, value in enumerate(values)
                    if all(k in value for k in include) and not any(k in value for k in exclude)), None)
   columns = {
       "destino": locate(["destino"]),
       "usuario": locate(["usuario"]),
       "tasa": locate(["tasa", "aplicable"]),
       "topeo": locate(["permite", "topeo"]),
       "prima": locate(["cobra", "prima"]),
       # Algunas sub-tablas usan "PRODUCTOS A UTILIZAR DESEMBOLSO"/"NUEVO PRODUCTO DESEMBOLSO",
       # otras solo "PRODUCTO "/"NUEVO PRODUCTO" (sin la palabra desembolso ni columna de
       # consolidación). "producto" a secas cubre ambas variantes; excluyendo "nuevo" y
       # "consolidacion" evitamos que una columna se cuele en la otra.
       "desembolso_actual": locate(["producto"], exclude=["nuevo", "consolidacion"]),
       "consolidacion_actual": locate(["producto", "consolidacion"], exclude=["nuevo"]),
       "desembolso_nuevo": locate(["nuevo", "producto"], exclude=["consolidacion"]),
       "consolidacion_nuevo": locate(["nuevo", "consolidacion"]),
       "cambio": locate(["cambio"]),
   }
   required = ("destino", "usuario", "desembolso_actual", "cambio")
   return columns if all(columns[key] is not None for key in required) else None

def scan_update_sheet(table: dict, regulation: str, effective_date: date, existing_products: set, existing_equivalences: dict, existing_conditions: dict | None = None) -> list:
   """Recorre una hoja de la planilla de actualización (una sub-tabla por Destino/Usuario,
   separadas por filas en blanco) y genera hallazgos usando el indicador ¿CAMBIÓ? y los
   pares PRODUCTOS A UTILIZAR (vigente) / NUEVO PRODUCTO (reemplazo) de la planilla.
   Además de Producto_Equivalencias (usada por "Conversión de productos"), busca la
   Condición_Reglamentacion vigente que usa el producto anterior -la que alimenta
   "Consulta normativa"- y propone su actualización con la misma vigencia, para que
   ambas pantallas queden sincronizadas."""
   existing_conditions = existing_conditions or {}
   location = table["location"]
   findings = []
   columns = None
   current_destino = current_usuario = ""
   for row_index, row in enumerate(table["rows"], start=1):
       if not any(clean(value) for value in row):
           columns, current_destino, current_usuario = None, "", ""
           continue
       candidate = update_sheet_columns(row)
       if candidate:
           columns, current_destino, current_usuario = candidate, "", ""
           continue
       if not columns:
           continue
       def cell(key):
           index = columns.get(key)
           return clean(row[index]) if index is not None and index < len(row) else ""
       current_destino = cell("destino") or current_destino
       current_usuario = cell("usuario") or current_usuario
       if not current_destino:
           continue
       evidence = " | ".join(clean(value) for value in row)
       cambio = cell("cambio").upper()
       desembolso_actual, consolidacion_actual = cell("desembolso_actual").upper(), cell("consolidacion_actual").upper()
       desembolso_nuevo, consolidacion_nuevo = cell("desembolso_nuevo").upper(), cell("consolidacion_nuevo").upper()
       if cambio != "SI":
           if desembolso_actual:
               findings.append(finding("EQUIVALENCIA", f"{location}, fila {row_index}", evidence, 90,
                   {"Destino": current_destino, "Codigo_Desembolso": desembolso_actual, "Codigo_Consolidacion": consolidacion_actual},
                   reason="La planilla indica ¿CAMBIÓ?=NO; se conserva como evidencia, no se genera una nueva versión."))
           continue
       nuevo_desembolso = desembolso_nuevo or desembolso_actual
       nuevo_consolidacion = consolidacion_nuevo or consolidacion_actual
       if not PRODUCT_PATTERN.fullmatch(nuevo_desembolso) or (nuevo_consolidacion and not PRODUCT_PATTERN.fullmatch(nuevo_consolidacion)):
           findings.append(finding("ADVERTENCIA", f"{location}, fila {row_index}", evidence, 30, {"Destino": current_destino},
               reason="¿CAMBIÓ?=SI pero no se pudo interpretar el nuevo código de producto; requiere revisión manual."))
           continue
       for code in (nuevo_desembolso, nuevo_consolidacion):
           if code and code not in existing_products:
               findings.append(finding("PRODUCTO", f"{location}, fila {row_index}", evidence, 99, {"Codigo": code},
                   "PRODUCTO", {"Accion": "GUARDAR", "Codigo": code, "Activo": "SI"}, code,
                   "Código nuevo detectado junto al reemplazo de producto de la planilla."))
               existing_products.add(code)
       # ------------------------------------------------------------------
       # CONDICION: es la tabla que realmente alimenta "Consulta normativa"
       # (Producto_Equivalencias sólo alimenta "Conversión de productos").
       # Se busca la condición vigente que hoy usa el producto anterior y se
       # propone renovarla con la misma Clave_Mantenimiento -para que el
       # publicado cierre la vieja por fecha y no la borre-, apuntando a los
       # productos nuevos y con la tasa/topeo que indique la planilla.
       # ------------------------------------------------------------------
       matches = existing_conditions.get(f"{desembolso_actual}|{regulation}", []) if desembolso_actual else []
       if not matches:
           findings.append(finding("ADVERTENCIA", f"{location}, fila {row_index}", evidence, 35,
               {"Destino": current_destino, "Codigo_Desembolso": desembolso_actual},
               reason=(f"¿CAMBIÓ?=SI pero no se encontró ninguna Condición vigente que use {desembolso_actual or '(sin código anterior)'} "
                       f"en {regulation}. La Conversión de productos va a quedar actualizada, pero la Consulta normativa no reflejará el "
                       "cambio salvo que se actualice la condición manualmente desde \"Cambio puntual\".")))
       else:
           tasa_nueva = cell("tasa")
           topeo_nuevo = cell("topeo")
           prima_nueva = cell("prima")
           for match in matches:
               proposal = {
                   "Accion": "GUARDAR",
                   "Clave_Mantenimiento": match["clave_mantenimiento"],
                   "Reglamentacion": regulation,
                   "Ambito_Geografico": match["ambito_geografico"],
                   "Destino": match["destino"],
                   "Tipo_Usuario": match["tipo_usuario"],
                   "Modalidad_Convenio": match["modalidad_convenio"],
                   "Fecha_Desde": effective_date,
                   "Fecha_Hasta": None,
                   "UVA_Desde": match["uva_desde"], "UVA_Desde_Inclusive": match["uva_desde_inclusive"],
                   "UVA_Hasta": match["uva_hasta"], "UVA_Hasta_Inclusive": match["uva_hasta_inclusive"],
                   "Tasa_Aplicable_Pct": tasa_nueva or match["tasa_aplicable_pct"],
                   "Permite_Topeo": topeo_nuevo or match["permite_topeo"],
                   "Cobra_Prima_Topeo": prima_nueva or match["cobra_prima_topeo"],
                   "Codigo_Producto": nuevo_desembolso if match["codigo_producto"] else None,
                   "Codigo_Desembolso": nuevo_desembolso if match["codigo_desembolso"] else None,
                   "Codigo_Consolidacion": nuevo_consolidacion if match["codigo_consolidacion"] else None,
                   "Referencia_Fuente": match["referencia_fuente"] or "Actualización automática (asistente)",
                   "Observaciones": f"Renovada por reemplazo de producto ({desembolso_actual}→{nuevo_desembolso}) detectado en {location}.",
               }
               findings.append(finding("CONDICION", f"{location}, fila {row_index}", evidence, 92,
                   {"Clave_Mantenimiento": match["clave_mantenimiento"], "Destino": match["destino"], "Codigo_Desembolso": desembolso_actual},
                   "CONDICION", proposal, f"{match['clave_mantenimiento']}|{effective_date.isoformat()}",
                   f"Renueva la condición vigente de {match['destino']} (antes con {desembolso_actual}) con el producto nuevo y la tasa de la planilla. "
                   "Revisar: la tasa/topeo se toma de la planilla; el resto de los campos se heredan de la condición anterior."))
       if not nuevo_consolidacion and not consolidacion_actual:
           # Esta sub-tabla no maneja producto de consolidación para este destino (no hay
           # columnas de consolidación). No corresponde crear/cerrar una equivalencia vacía;
           # sólo se registra el alta del producto de desembolso nuevo y de la condición, ya hechas arriba.
           continue
       # Si el desembolso nuevo es un código distinto del vigente, el par anterior no queda
       # "enganchado" automáticamente al publicar (la publicación cierra por Producto_Desembolso_ID
       # igual al nuevo), así que acá se propone explícitamente su baja para no perder el historial.
       replaces_pair = desembolso_actual and desembolso_actual != nuevo_desembolso
       if replaces_pair:
           findings.append(finding("EQUIVALENCIA", f"{location}, fila {row_index}", evidence, 95,
               {"Codigo_Desembolso": desembolso_actual, "Codigo_Consolidacion": consolidacion_actual},
               "EQUIVALENCIA", {"Accion": "BAJA", "Codigo_Desembolso": desembolso_actual, "Reglamentacion": regulation,
               "Fecha_Hasta": effective_date - timedelta(days=1)},
               f"{desembolso_actual}|{regulation}|BAJA|{effective_date.isoformat()}",
               f"¿CAMBIÓ?=SI: {desembolso_actual}→{consolidacion_actual} es reemplazado por {nuevo_desembolso}→{nuevo_consolidacion}. "
               "Se cierra su vigencia; no se elimina, queda en el historial."))
       active_consolidation = existing_equivalences.get(f"{nuevo_desembolso}|{regulation}")
       if not replaces_pair and active_consolidation == nuevo_consolidacion:
           findings.append(finding("EQUIVALENCIA", f"{location}, fila {row_index}", evidence, 90,
               {"Codigo_Desembolso": nuevo_desembolso, "Codigo_Consolidacion": nuevo_consolidacion},
               reason=f"{nuevo_desembolso} ya tiene a {nuevo_consolidacion} como consolidación vigente en {regulation}; no se generará una nueva versión."))
           continue
       reason = f"¿CAMBIÓ?=SI en la planilla para {current_destino}."
       if replaces_pair:
           reason += f" Alta del nuevo par {nuevo_desembolso}→{nuevo_consolidacion}, junto con la baja del anterior."
       elif active_consolidation:
           reason += f" Reemplaza a la consolidación anterior ({active_consolidation}) del mismo producto, vigente hasta el día previo."
       findings.append(finding("EQUIVALENCIA", f"{location}, fila {row_index}", evidence, 97,
           {"Codigo_Desembolso": nuevo_desembolso, "Codigo_Consolidacion": nuevo_consolidacion},
           "EQUIVALENCIA", {"Accion": "GUARDAR", "Codigo_Desembolso": nuevo_desembolso, "Codigo_Consolidacion": nuevo_consolidacion,
           "Reglamentacion": regulation, "Fecha_Desde": effective_date, "Fecha_Hasta": None},
           f"{nuevo_desembolso}|{regulation}|{effective_date.isoformat()}", reason))
       existing_equivalences[f"{nuevo_desembolso}|{regulation}"] = nuevo_consolidacion
   return findings

def organization_columns(headers):
   values = [normalized(value) for value in headers]
   def locate(predicate):
       return next((index for index, value in enumerate(values) if predicate(value)), None)
   columns = {
       "name": locate(lambda value: "organismo" in value),
       "cuit": locate(lambda value: "cuit" in value),
       "first": locate(lambda value: "tasa" in value and any(token in value for token in ("1ra", "primera", "1 vivienda"))),
       "second": locate(lambda value: "tasa" in value and any(token in value for token in ("2da", "segunda", "2 vivienda"))),
       "until": locate(lambda value: "vigencia" in value or ("recepcion" in value and "solicitud" in value)),
   }
   return columns if columns["name"] is not None and columns["cuit"] is not None and columns["first"] is not None else None

def product_columns(headers):
   values = [normalized(value) for value in headers]
   disbursement = next((i for i, value in enumerate(values) if "desembolso" in value), None)
   consolidation = next((i for i, value in enumerate(values) if "consolidacion" in value), None)
   return (disbursement, consolidation) if disbursement is not None and consolidation is not None else None

def match_organism(name, cuit, existing, document_id, sequence):
   name_key = normalized(name)
   exact = [item for item in existing if clean(item.get("cuit")) == cuit and normalized(item.get("nombre")) == name_key]
   if exact:
       return exact[0]["codigo"], 100, True
   candidates = [item for item in existing if clean(item.get("cuit")) == cuit]
   if candidates:
       best = max(candidates, key=lambda item: SequenceMatcher(None, normalized(item.get("nombre")), name_key).ratio())
       similarity = SequenceMatcher(None, normalized(best.get("nombre")), name_key).ratio()
       if similarity >= .80:
           return best["codigo"], 92 + similarity * 7, True
   return f"D{document_id % 100000:05d}O{sequence:04d}", 96, False

def table_policy(table_index, table_count, all_text):
   has_pre = "1 5 p p a" in normalized(all_text) or "adicional de 1 5" in normalized(all_text)
   has_post = "2 p p a" in normalized(all_text) or "adicional de 2" in normalized(all_text)
   if table_count >= 2 and has_pre and has_post:
       return ("PRE_3214", Decimal("1.50"), False) if table_index == 1 else ("POST_3214", Decimal("2.00"), True)
   if "circular n 3214" in normalized(all_text) or "circular 3214" in normalized(all_text):
       return "POST_3214", Decimal("2.00"), True
   return "PRE_3214", Decimal("1.50"), False

def analyze(content: bytes, extension: str, document_id: int, effective_date: date, existing_organisms=None, existing_products=None, existing_equivalences=None, existing_conditions=None):
   existing_organisms = existing_organisms or []
   existing_products = {clean(value).upper() for value in (existing_products or [])}
   # existing_equivalences: dict "CODIGO_DESEMBOLSO|REGLAMENTACION" -> CODIGO_CONSOLIDACION vigente hoy.
   existing_equivalences = existing_equivalences or {}
   # existing_conditions: dict "CODIGO_PRODUCTO_ANCLA|REGLAMENTACION" -> lista de condiciones vigentes
   # (una por cada Destino/Ámbito/Usuario que use ese producto) para poder renovarlas.
   existing_conditions = existing_conditions or {}
   if extension == "docx":
       source = extract_docx(content)
   elif extension == "pdf":
       source = extract_pdf(content)
   elif extension == "xlsx":
       source = extract_xlsx(content)
   else:
       raise ValueError("Formato no soportado")
   all_text = "\n".join(source["paragraphs"] + [" | ".join(row) for table in source["tables"] for row in table["rows"]])
   regulation, regulation_confidence = detect_regulation(all_text)
   findings = []
   if regulation:
       findings.append(finding("REGLAMENTACION", "Documento", regulation, regulation_confidence, {"Reglamentacion": regulation}, reason="Detectada por título y contenido."))
   else:
       findings.append(finding("ADVERTENCIA", "Documento", "No se pudo determinar la reglamentación.", 20, {}, reason="Seleccione o indique la reglamentación antes de generar una carga."))
   for warning in source["warnings"]:
       findings.append(finding("ADVERTENCIA", "Documento", warning, 20, {}, reason=warning))
   organization_sequence = 0
   proposed_organisms = set()
   handled_evidence = set()
   for table_index, table in enumerate(source["tables"], start=1):
       if not table["rows"]:
           continue
       sheet_reg = sheet_regulation(table["location"]) if extension == "xlsx" else None
       if sheet_reg:
           findings.extend(scan_update_sheet(table, sheet_reg, effective_date, existing_products, existing_equivalences, existing_conditions))
           handled_evidence.update(" | ".join(clean(value) for value in row) for row in table["rows"])
           continue
       headers = table["rows"][0]
       org_columns = organization_columns(headers)
       if org_columns:
           group, additional, applies_circular = table_policy(table_index, len(source["tables"]), all_text)
           for row_index, row in enumerate(table["rows"][1:], start=2):
               def cell(index): return row[index] if index is not None and index < len(row) else ""
               name, cuit = cell(org_columns["name"]), re.sub(r"\D", "", cell(org_columns["cuit"]))
               if not name or len(cuit) != 11:
                   continue
               organization_sequence += 1
               code, match_confidence, exists = match_organism(name, cuit, existing_organisms, document_id, organization_sequence)
               evidence = " | ".join(row)
               org_proposal = {"Accion": "GUARDAR", "Codigo": code, "Nombre": name, "CUIT": cuit, "Activo": "SI"}
               proposal_type = None if code in proposed_organisms else "ORGANISMO"
               findings.append(finding("ORGANISMO", f"{table['location']}, fila {row_index}", evidence, match_confidence, {"Nombre": name, "CUIT": cuit}, proposal_type, org_proposal if proposal_type else None, code, "Organismo existente reconocido." if exists else "Código sugerido automáticamente para un organismo nuevo."))
               proposed_organisms.add(code)
               if not exists:
                   existing_organisms.append({"codigo": code, "nombre": name, "cuit": cuit})
               until = parse_date(cell(org_columns["until"]))
               for dwelling, column in (("UNICA", org_columns["first"]), ("SEGUNDA", org_columns["second"])):
                   raw_rate = cell(column)
                   rate = parse_rate(raw_rate)
                   if rate is None:
                       if dwelling == "SEGUNDA" and ("---" in raw_rate or not clean(raw_rate)):
                           continue
                       findings.append(finding("ORGANISMO_TASA", f"{table['location']}, fila {row_index}", evidence, 55, {"Codigo_Organismo": code, "Tipo_Vivienda": dwelling, "Tasa": raw_rate}, reason="No se pudo interpretar la tasa."))
                       continue
                   valid_period = until is not None and until >= effective_date
                   proposed = None
                   confidence = 97 if valid_period else 72
                   reason = "Tasa, vivienda y vigencia identificadas en columnas explícitas."
                   if valid_period:
                       proposed = {
                           "Accion": "GUARDAR", "Codigo_Organismo": code, "Tipo_Vivienda": dwelling,
                           "Tasa_Pct": rate, "Grupo_Pauta": group, "Vigencia_Desde": effective_date,
                           "Vigencia_Hasta": until, "Adicional_Topeo_Pct": additional if dwelling == "UNICA" else None,
                           "Aplica_Circular_3214": "SI" if applies_circular else "NO",
                           "Observaciones": "Extraído automáticamente; fecha desde indicada por el usuario durante el análisis.",
                       }
                   else:
                       reason = "La vigencia está vencida o falta la fecha desde; se conserva como evidencia, sin propuesta publicable."
                   findings.append(finding("ORGANISMO_TASA", f"{table['location']}, fila {row_index}", evidence, confidence, {"Codigo_Organismo": code, "Tipo_Vivienda": dwelling, "Tasa_Pct": rate, "Vigencia_Hasta": until}, "ORGANISMO_TASA" if proposed else None, proposed, f"{code}|{dwelling}|{effective_date.isoformat()}", reason))
               handled_evidence.add(evidence)
           continue
       product_pair = product_columns(headers)
       if product_pair:
           for row_index, row in enumerate(table["rows"][1:], start=2):
               disbursement = clean(row[product_pair[0]]).upper() if product_pair[0] < len(row) else ""
               consolidation = clean(row[product_pair[1]]).upper() if product_pair[1] < len(row) else ""
               if not PRODUCT_PATTERN.fullmatch(disbursement) or not PRODUCT_PATTERN.fullmatch(consolidation):
                   continue
               evidence = " | ".join(row)
               for code in (disbursement, consolidation):
                   if code not in existing_products:
                       findings.append(finding("PRODUCTO", f"{table['location']}, fila {row_index}", evidence, 99, {"Codigo": code}, "PRODUCTO", {"Accion": "GUARDAR", "Codigo": code, "Activo": "SI"}, code, "Código nuevo detectado en una columna de producto."))
                       existing_products.add(code)
               if regulation:
                   active_key = f"{disbursement}|{regulation}"
                   active_consolidation = existing_equivalences.get(active_key)
                   if active_consolidation == consolidation:
                       # Misma consolidación ya vigente: es el mismo Excel recurrente sin novedades
                       # para este producto. Se deja como evidencia, sin generar una nueva versión.
                       findings.append(finding(
                           "EQUIVALENCIA", f"{table['location']}, fila {row_index}", evidence, 90,
                           {"Codigo_Desembolso": disbursement, "Codigo_Consolidacion": consolidation},
                           reason=f"{disbursement} ya tiene a {consolidation} como consolidación vigente en {regulation}; no se generará una nueva versión.",
                       ))
                   else:
                       reason = "Columnas de desembolso y consolidación identificadas explícitamente."
                       if active_consolidation:
                           reason += (f" Reemplaza a la consolidación anterior ({active_consolidation}), que quedará "
                                      "vigente hasta el día previo a la nueva fecha desde; no se elimina, se cierra su historial.")
                       else:
                           reason += " No hay una equivalencia vigente previa para este producto en esta reglamentación."
                       proposal = {"Accion": "GUARDAR", "Codigo_Desembolso": disbursement, "Codigo_Consolidacion": consolidation, "Reglamentacion": regulation, "Fecha_Desde": effective_date, "Fecha_Hasta": None}
                       findings.append(finding("EQUIVALENCIA", f"{table['location']}, fila {row_index}", evidence, 96, {"Codigo_Desembolso": disbursement, "Codigo_Consolidacion": consolidation}, "EQUIVALENCIA", proposal, f"{disbursement}|{regulation}|{effective_date.isoformat()}", reason))
               handled_evidence.add(evidence)
   # Descubrimiento genérico: informa tasas y productos que no pudieron convertirse con seguridad.
   generic_sources = [(f"Párrafo {index}", text) for index, text in enumerate(source["paragraphs"], start=1)]
   for table in source["tables"]:
       generic_sources.extend((f"{table['location']}, fila {index}", " | ".join(row)) for index, row in enumerate(table["rows"], start=1))
   seen_generic = set()
   for location, evidence in generic_sources:
       if evidence in handled_evidence:
           continue
       for code in PRODUCT_PATTERN.findall(evidence):
           code = code.upper()
           key = ("PRODUCTO", code, location)
           if key in seen_generic: continue
           seen_generic.add(key)
           proposed = None if code in existing_products else {"Accion": "GUARDAR", "Codigo": code, "Activo": "SI"}
           findings.append(finding("PRODUCTO", location, evidence, 91 if proposed else 84, {"Codigo": code}, "PRODUCTO" if proposed else None, proposed, code, "Producto detectado por patrón; revise su función en el cuadro."))
       for raw_rate in RATE_PATTERN.findall(evidence):
           key = ("TASA", raw_rate, location)
           if key in seen_generic: continue
           seen_generic.add(key)
           findings.append(finding("TASA", location, evidence, 68, {"Tasa_Pct": Decimal(raw_rate.replace(",", "."))}, reason="Se detectó la tasa, pero faltan dimensiones suficientes para crear una condición automáticamente."))
       if len(findings) >= 1500:
           findings.append(finding("ADVERTENCIA", "Documento", "El análisis alcanzó el máximo de 1.500 hallazgos.", 40, {}, reason="Divida el documento o utilice la plantilla masiva."))
           break
   # Una clave solo puede originar una modificación. Las repeticiones idénticas se
   # conservan como evidencia; los valores contradictorios se bloquean para revisión.
   proposal_by_key = {}
   for item in findings:
       identity = (item["proposal_type"], item["proposal_key"])
       if not all(identity):
           continue
       previous = proposal_by_key.get(identity)
       if previous is None:
           proposal_by_key[identity] = item
           continue
       if previous["proposed"] == item["proposed"]:
           item["proposal_type"] = None
           item["proposed"] = None
           item["status"] = "DESCARTADO"
           item["selected"] = False
           item["reason"] = "Repetición idéntica conservada únicamente como evidencia."
       else:
           for conflicting in (previous, item):
               conflicting["proposal_type"] = None
               conflicting["status"] = "REQUIERE_REVISION"
               conflicting["selected"] = False
               conflicting["reason"] = "La misma clave contiene valores contradictorios; debe resolverse manualmente."
   proposals = [item for item in findings if item["proposal_type"]]
   review = [item for item in findings if item["status"] == "REQUIERE_REVISION"]
   weighted = sum(item["confidence"] for item in proposals) / len(proposals) if proposals else regulation_confidence
   return {
       "regulation": regulation,
       "confidence": round(weighted, 2),
       "findings": findings,
       "total_findings": len(findings),
       "total_proposals": len(proposals),
       "total_reviews": len(review),
   }